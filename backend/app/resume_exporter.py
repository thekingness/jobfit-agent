import json
import re
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


def _safe_list(value: Any) -> List[Any]:
    return value if isinstance(value, list) else []


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _normalize_text(text: str) -> str:
    text = _safe_text(text)
    text = re.sub(r"\s+", "", text)
    text = text.replace("•", "").replace("·", "")
    return text


def _iter_paragraphs_in_document(document: DocxDocument):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for paragraph in nested_cell.paragraphs:
                                yield paragraph


def _replace_paragraph_keep_style(paragraph, new_text: str) -> None:
    """
    尽量保留段落级格式、项目符号、表格位置和第一个 run 的字体样式。
    不直接重建文档，因此原始 DOCX 里的图片、页眉、页脚、表格布局通常都会保留。
    """
    new_text = _safe_text(new_text)

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _try_replace_in_single_run(paragraph, original: str, optimized: str) -> bool:
    for run in paragraph.runs:
        if original and original in run.text:
            run.text = run.text.replace(original, optimized)
            return True
    return False


def _paragraph_should_replace(paragraph_text: str, original: str) -> bool:
    paragraph_text = _safe_text(paragraph_text)
    original = _safe_text(original)

    if not paragraph_text or not original:
        return False

    if original in paragraph_text:
        return True

    norm_para = _normalize_text(paragraph_text)
    norm_original = _normalize_text(original)

    if not norm_para or not norm_original:
        return False

    if norm_original in norm_para:
        return True

    # 原文较长时，允许“解析结果”和 Word 段落有轻微差异。
    if len(norm_original) >= 16:
        overlap = len(set(norm_para) & set(norm_original)) / max(len(set(norm_original)), 1)
        length_ratio = min(len(norm_para), len(norm_original)) / max(len(norm_para), len(norm_original), 1)
        if overlap >= 0.72 and length_ratio >= 0.45:
            return True

    return False


def _build_replacement_pairs(payload: Dict[str, Any]) -> List[Tuple[str, str]]:
    report_json = payload.get("report_json", {}) or {}
    rewrite_suggestions = _safe_list(report_json.get("rewrite_suggestions"))

    pairs: List[Tuple[str, str]] = []
    seen = set()

    for item in rewrite_suggestions:
        if not isinstance(item, dict):
            continue

        original = _safe_text(item.get("original"))
        optimized = _safe_text(item.get("optimized"))
        risk_level = _safe_text(item.get("risk_level"))

        # 高风险改写可能缺少原文证据，默认不自动写入简历，避免伪造经历。
        if not original or not optimized or "高" in risk_level:
            continue

        key = (_normalize_text(original), _normalize_text(optimized))
        if key in seen:
            continue

        seen.add(key)
        pairs.append((original, optimized))

    return pairs


def build_preserve_format_docx(original_docx_path: str, payload: Dict[str, Any]) -> BytesIO:
    """
    原格式优化版：在原始 DOCX 上定点替换文本。

    优点：
    - 原照片通常保留
    - 原表格/页眉页脚/版式通常保留
    - 只替换 report_json.rewrite_suggestions 中的低/中风险改写内容

    限制：
    - 只支持 DOCX
    - 如果模型返回的 original 和原文差异太大，可能无法定位该段落
    """
    document = Document(original_docx_path)
    pairs = _build_replacement_pairs(payload)

    used_paragraph_ids = set()

    for original, optimized in pairs:
        replaced = False

        for paragraph in _iter_paragraphs_in_document(document):
            para_id = id(paragraph._p)
            if para_id in used_paragraph_ids:
                continue

            para_text = paragraph.text or ""

            if not _paragraph_should_replace(para_text, original):
                continue

            if original in para_text and _try_replace_in_single_run(paragraph, original, optimized):
                pass
            else:
                _replace_paragraph_keep_style(paragraph, optimized)

            used_paragraph_ids.add(para_id)
            replaced = True
            break

        # 定位不到时不强行添加，避免破坏原模板。
        if not replaced:
            continue

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_vertical_center(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def _set_document_base_style(document: DocxDocument) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)

    for section in document.sections:
        section.top_margin = Cm(1.35)
        section.bottom_margin = Cm(1.35)
        section.left_margin = Cm(1.45)
        section.right_margin = Cm(1.45)


def _extract_first_image_from_docx(docx_path: str) -> Optional[BytesIO]:
    path = Path(docx_path)
    if path.suffix.lower() != ".docx":
        return None

    try:
        with zipfile.ZipFile(docx_path, "r") as archive:
            media_files = [
                name for name in archive.namelist()
                if name.startswith("word/media/")
                and name.lower().endswith((".png", ".jpg", ".jpeg"))
            ]
            if not media_files:
                return None

            image_bytes = archive.read(media_files[0])
            return BytesIO(image_bytes)
    except Exception:
        return None


def _add_header(document: DocxDocument, payload: Dict[str, Any], image_stream: Optional[BytesIO]) -> None:
    resume_info = payload.get("resume_info", {}) or {}
    jd_info = payload.get("jd_info", {}) or {}
    match_result = payload.get("match_result", {}) or {}

    name = _safe_text(resume_info.get("name")) or "候选人"
    position = _safe_text(jd_info.get("position")) or "目标岗位"
    score = match_result.get("score", "-")

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    _set_cell_shading(left, "EFF6FF")
    _set_cell_shading(right, "EFF6FF")
    _set_cell_vertical_center(left)
    _set_cell_vertical_center(right)

    p = left.paragraphs[0]
    _set_paragraph_spacing(p, after=4)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(15, 23, 42)

    p2 = left.add_paragraph()
    _set_paragraph_spacing(p2, after=3)
    r2 = p2.add_run(f"求职目标：{position}")
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(30, 64, 175)

    p3 = left.add_paragraph()
    r3 = p3.add_run(f"岗位匹配分：{score}")
    r3.font.size = Pt(10.5)
    r3.font.color.rgb = RGBColor(71, 85, 105)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_stream:
        try:
            rp.add_run().add_picture(image_stream, width=Cm(2.7))
        except Exception:
            rp.add_run("照片")
    else:
        run = rp.add_run("照片区域")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 116, 139)

    document.add_paragraph()


def _add_section_title(document: DocxDocument, title: str) -> None:
    p = document.add_paragraph()
    _set_paragraph_spacing(p, before=7, after=3)

    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(30, 64, 175)

    # 添加简洁分隔线
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "BFDBFE")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def _add_bullet(document: DocxDocument, text: str) -> None:
    text = _safe_text(text)
    if not text:
        return

    p = document.add_paragraph()
    _set_paragraph_spacing(p, after=2)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.25)

    run = p.add_run("• " + text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(51, 65, 85)


def _add_plain_line(document: DocxDocument, text: str, bold: bool = False) -> None:
    p = document.add_paragraph()
    _set_paragraph_spacing(p, after=2)
    run = p.add_run(_safe_text(text))
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(51, 65, 85)


def build_beautified_resume_docx(original_file_path: str, payload: Dict[str, Any]) -> BytesIO:
    """
    模板美化版：重新排版生成一份更美观的 DOCX。
    会尝试从原 DOCX 中提取第一张图片作为头像；PDF/TXT 暂不保证提取照片。
    """
    resume_info = payload.get("resume_info", {}) or {}
    jd_info = payload.get("jd_info", {}) or {}
    report_json = payload.get("report_json", {}) or {}
    match_result = payload.get("match_result", {}) or {}

    document = Document()
    _set_document_base_style(document)

    image_stream = _extract_first_image_from_docx(original_file_path)
    _add_header(document, payload, image_stream)

    advantages = _safe_list(report_json.get("advantages"))
    skills = _safe_list(resume_info.get("skills"))
    education = _safe_list(resume_info.get("education"))
    internships = _safe_list(resume_info.get("internships"))
    certificates = _safe_list(resume_info.get("certificates"))
    projects = _safe_list(resume_info.get("projects"))
    rewrite_suggestions = _safe_list(report_json.get("rewrite_suggestions"))
    matched_skills = _safe_list(match_result.get("matched_skills"))

    _add_section_title(document, "个人优势")
    if advantages:
        for item in advantages[:5]:
            _add_bullet(document, str(item))
    else:
        _add_bullet(document, "具备与目标岗位相关的技术基础和项目经历，可继续补充量化成果。")

    _add_section_title(document, "技能栈")
    merged_skills: List[str] = []
    for item in matched_skills + skills:
        text = _safe_text(item)
        if text and text not in merged_skills:
            merged_skills.append(text)

    if merged_skills:
        _add_plain_line(document, "、".join(merged_skills))
    else:
        _add_bullet(document, "未解析到明确技能，请人工补充专业技能。")

    _add_section_title(document, "教育背景")
    if education:
        for item in education:
            _add_bullet(document, str(item))
    else:
        _add_bullet(document, "暂无教育背景信息，请根据真实情况补充。")

    _add_section_title(document, "项目经历")
    usable_rewrites = []
    for item in rewrite_suggestions:
        if not isinstance(item, dict):
            continue
        optimized = _safe_text(item.get("optimized"))
        risk_level = _safe_text(item.get("risk_level")) or "低"
        evidence = _safe_text(item.get("evidence"))
        if optimized and "高" not in risk_level:
            usable_rewrites.append((optimized, evidence, risk_level))

    if usable_rewrites:
        for optimized, evidence, risk_level in usable_rewrites:
            _add_bullet(document, optimized)
            if evidence and "中" in risk_level:
                _add_bullet(document, f"说明：该条建议基于“{evidence}”，建议投递前人工确认。")
    elif projects:
        for project in projects:
            if not isinstance(project, dict):
                _add_bullet(document, str(project))
                continue

            name = _safe_text(project.get("name")) or "项目经历"
            tech_stack = _safe_list(project.get("tech_stack"))
            responsibilities = _safe_list(project.get("responsibilities"))
            achievements = _safe_list(project.get("achievements"))

            _add_plain_line(
                document,
                f"{name}" + (f"｜技术栈：{'、'.join(map(str, tech_stack))}" if tech_stack else ""),
                bold=True,
            )
            for item in responsibilities + achievements:
                _add_bullet(document, str(item))
    else:
        _add_bullet(document, "暂无项目经历信息，请根据真实项目补充。")

    if internships:
        _add_section_title(document, "实习 / 工作经历")
        for item in internships:
            _add_bullet(document, str(item))

    if certificates:
        _add_section_title(document, "证书 / 荣誉")
        for item in certificates:
            _add_bullet(document, str(item))

    risk_tips = _safe_list(report_json.get("risk_tips"))
    _add_section_title(document, "投递前检查")
    if risk_tips:
        for item in risk_tips[:3]:
            _add_bullet(document, str(item))
    _add_bullet(document, "本简历由 JobFit Agent 根据原简历和岗位 JD 生成，投递前请人工核对经历、时间、数据和技术表述。")

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
